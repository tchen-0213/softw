#!/usr/bin/env python3
"""Build formal Word/PDF submission copies with embedded figures."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

import markdown
import pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[2]
IMG_DIR = ROOT / "02_docs" / "images"
PNG_DIR = IMG_DIR / "png"
OUT_DIR = ROOT / "提交包"
CHROME = "/opt/google/chrome/chrome"
FONT_ASIAN = "Microsoft YaHei"

DOCUMENTS = [
    (ROOT / "02_docs" / "软件需求规格说明书.md", "02-需求规格说明书"),
    (ROOT / "02_docs" / "软件概要设计说明书.md", "02-概要设计说明书"),
    (ROOT / "02_docs" / "软件详细设计说明书.md", "02-详细设计说明书"),
    (ROOT / "02_docs" / "业务场景用例清单与追溯表.md", "02-业务场景用例清单与追溯表"),
    (ROOT / "02_docs" / "第1天业务场景清单与确认表.md", "02-第1天业务场景清单与确认表"),
    (ROOT / "02_docs" / "测试文档.md", "02-测试计划"),
    (ROOT / "04_tests" / "reports" / "tests" / "测试报告-小学期.md", "04-测试报告"),
    (ROOT / "02_docs" / "微服务拆分设计.md", "02-微服务拆分设计"),
    (ROOT / "02_docs" / "微服务接口与数据归属.md", "02-微服务接口与数据归属"),
    (ROOT / "04_tests" / "reports" / "performance" / "性能对比实验报告.md", "04-性能对比实验报告"),
    (ROOT / "06_defense" / "技术总结报告.md", "06-技术总结报告"),
    (ROOT / "06_defense" / "最终交付核查清单.md", "06-最终交付核查清单"),
    (ROOT / "06_defense" / "答辩提纲.md", "06-答辩提纲"),
    (ROOT / "05_management" / "个人权重表.md", "05-个人权重表"),
    (ROOT / "05_management" / "全员确认记录.md", "05-全员确认记录"),
]

PDF_MIRROR = {
    "02-需求规格说明书": ROOT / "02_docs" / "pdf" / "软件需求规格说明书.pdf",
    "02-概要设计说明书": ROOT / "02_docs" / "pdf" / "软件概要设计说明书.pdf",
    "02-详细设计说明书": ROOT / "02_docs" / "pdf" / "软件详细设计说明书.pdf",
    "02-业务场景用例清单与追溯表": ROOT / "02_docs" / "pdf" / "业务场景用例清单与追溯表.pdf",
    "02-测试计划": ROOT / "02_docs" / "pdf" / "测试文档.pdf",
    "04-测试报告": ROOT / "02_docs" / "pdf" / "测试报告-小学期.pdf",
    "02-微服务接口与数据归属": ROOT / "02_docs" / "pdf" / "微服务接口与数据归属.pdf",
    "04-性能对比实验报告": ROOT / "02_docs" / "pdf" / "性能对比实验报告.pdf",
    "06-技术总结报告": ROOT / "02_docs" / "pdf" / "技术总结报告.pdf",
    "02-微服务拆分设计": ROOT / "02_docs" / "pdf" / "微服务拆分设计.pdf",
    "02-第1天业务场景清单与确认表": ROOT / "02_docs" / "pdf" / "第1天业务场景清单与确认表.pdf",
    "06-最终交付核查清单": ROOT / "02_docs" / "pdf" / "最终交付核查清单.pdf",
    "06-答辩提纲": ROOT / "02_docs" / "pdf" / "答辩提纲.pdf",
}

HTML_CSS = """
body { font-family: "Noto Sans CJK SC", "WenQuanYi Micro Hei", "Microsoft YaHei", sans-serif;
       font-size: 14px; line-height: 1.6; color: #1f2933; margin: 18mm; }
h1 { font-size: 22px; border-bottom: 2px solid #1d4ed8; padding-bottom: 8px; }
h2 { font-size: 18px; margin-top: 28px; }
h3 { font-size: 16px; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 12px; }
th, td { border: 1px solid #cbd5e1; padding: 6px 8px; vertical-align: top; }
th { background: #eef2ff; }
img { max-width: 100%; height: auto; display: block; margin: 10px auto; }
pre, code { font-family: ui-monospace, Consolas, monospace; font-size: 12px; }
pre { background: #f8fafc; padding: 10px; white-space: pre-wrap; }
blockquote { color: #475569; border-left: 4px solid #93c5fd; padding-left: 10px; }
"""


def svg_to_png() -> None:
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    for svg in sorted(IMG_DIR.glob("*.svg")):
        png = PNG_DIR / f"{svg.stem}.png"
        if png.exists() and png.stat().st_mtime >= svg.stat().st_mtime:
            continue
        doc = pymupdf.open(svg)
        pix = doc[0].get_pixmap(dpi=160, alpha=False)
        pix.save(png)


def rewrite_images(html: str, base: Path) -> str:
    def repl(match: re.Match[str]) -> str:
        src = match.group(1)
        if src.startswith("http"):
            return match.group(0)
        src_path = (base / src).resolve()
        if src_path.suffix.lower() == ".svg":
            png_path = PNG_DIR / f"{src_path.stem}.png"
            if png_path.exists():
                src_path = png_path
        return f'<img src="{src_path.as_uri()}" alt="" style="max-width:100%">'

    return re.sub(r'<img[^>]*src="([^"]+)"[^>]*>', repl, html)


def md_to_html(md_path: Path, title: str) -> str:
    body = markdown.markdown(
        md_path.read_text(encoding="utf-8"),
        extensions=["extra", "tables", "sane_lists", "nl2br"],
    )
    body = rewrite_images(body, md_path.parent)
    return (
        "<!DOCTYPE html><html lang='zh-CN'><head><meta charset='utf-8'>"
        f"<title>{title}</title><style>{HTML_CSS}</style></head>"
        f"<body><h1>{title}</h1>{body}</body></html>"
    )


def html_to_pdf(html_path: Path, pdf_path: Path) -> None:
    profile = Path("/tmp") / f"chrome-submit-{html_path.stem}"
    profile.mkdir(parents=True, exist_ok=True)
    cmd = [
        CHROME, "--headless=new", "--no-sandbox", "--disable-gpu",
        "--disable-dev-shm-usage", "--no-pdf-header-footer",
        f"--user-data-dir={profile}",
        f"--print-to-pdf={pdf_path}",
        html_path.resolve().as_uri(),
    ]
    result = subprocess.run(cmd, check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
    if result.returncode != 0 or not pdf_path.exists() or pdf_path.stat().st_size < 2000:
        raise RuntimeError(f"PDF failed for {html_path.name}: {result.stderr[-800:]}")


def set_run_font(run) -> None:
    run.font.name = FONT_ASIAN
    run.font.size = Pt(11)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), FONT_ASIAN)


def add_text(paragraph, text: str, bold: bool = False, size: int = 11) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_ASIAN
    rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), FONT_ASIAN)


def md_to_docx(md_path: Path, docx_path: Path, title: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    heading = doc.add_heading(title, 0)
    for run in heading.runs:
        set_run_font(run)
        run.font.size = Pt(18)

    lines = md_path.read_text(encoding="utf-8").replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code: list[str] = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                add_text(p, "\n".join(code), size=9)
                code = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code.append(line)
            i += 1
            continue
        image = re.search(r"!\[[^\]]*\]\(([^)]+)\)", line)
        if image:
            src = (md_path.parent / image.group(1)).resolve()
            png = PNG_DIR / f"{src.stem}.png" if src.suffix.lower() == ".svg" else src
            if png.exists():
                doc.add_picture(str(png), width=Cm(15.5))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(cap, Path(image.group(1)).name, size=9)
            i += 1
            continue
        if re.match(r"^\s*\|.*\|\s*$", line):
            rows = []
            while i < len(lines) and re.match(r"^\s*\|.*\|\s*$", lines[i]):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.match(r"^:?-+:?$", c or "") for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = row[c_idx] if c_idx < len(row) else ""
                        for p in cell.paragraphs:
                            for run in p.runs:
                                set_run_font(run)
                                run.font.size = Pt(9)
            continue
        heading_m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_m:
            level = min(len(heading_m.group(1)), 4)
            h = doc.add_heading(heading_m.group(2).strip(), level)
            for run in h.runs:
                set_run_font(run)
            i += 1
            continue
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line).strip()
        if text:
            p = doc.add_paragraph()
            add_text(p, text)
        i += 1
    doc.save(docx_path)


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    svg_to_png()
    html_dir = Path("/tmp") / "moyu-submit-html"
    html_dir.mkdir(exist_ok=True)
    for md_path, stem in DOCUMENTS:
        if not md_path.exists():
            print(f"missing {md_path}", file=sys.stderr)
            return 1
        title = md_path.stem
        html_path = html_dir / f"{stem}.html"
        pdf_path = OUT_DIR / f"{stem}.pdf"
        docx_path = OUT_DIR / f"{stem}.docx"
        html_path.write_text(md_to_html(md_path, title), encoding="utf-8")
        html_to_pdf(html_path, pdf_path)
        md_to_docx(md_path, docx_path, title)
        mirror = PDF_MIRROR.get(stem)
        if mirror is not None:
            mirror.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(pdf_path, mirror)
        print(f"{stem}.pdf {pdf_path.stat().st_size}  {stem}.docx {docx_path.stat().st_size}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
